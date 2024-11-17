import pandas as pd
import yaml
from collections import Counter
import plotly.graph_objects as go
import re
import pandas as pd
import PyPDF2
import plotly.express as px
from docx import Document
from googletrans import Translator
from concurrent.futures import ThreadPoolExecutor, as_completed
from docx import Document
from docx.shared import Pt

def grapth(yaml_file, bunch_file, output_file):
    # Загрузка названий столбцов из YAML
    with open(yaml_file, 'r', encoding='utf-8') as file:
        columns = yaml.safe_load(file)

    # Загрузка данных из Excel
    df = pd.read_excel(bunch_file, sheet_name='SHEET')

    # Функция для извлечения количества по годам
    def extract_year_counts(df, column_name):
        valid_entries = df[column_name].dropna()
        years = valid_entries[valid_entries != 'no data'].str[:4]
        year_counts = Counter(years)
        return pd.DataFrame(list(year_counts.items()), columns=['Year', column_name.split()[1].capitalize()])

    # Извлечение данных для каждого столбца
    priority_df = extract_year_counts(df, columns['priority_date_column'])
    publication_df = extract_year_counts(df, columns['publication_date_column'])
    grant_df = extract_year_counts(df, columns['grant_date_column'])

    # Объединение датафреймов
    merged_df = priority_df.merge(publication_df, on='Year', how='outer').merge(grant_df, on='Year', how='outer').fillna(0)

    # Фильтрация данных по годам от 2014 до 2023
    merged_df['Year'] = merged_df['Year'].astype(int)
    merged_df = merged_df[(merged_df['Year'] >= 2014) & (merged_df['Year'] <= 2023)]
    merged_df = merged_df.sort_values(by='Year')

    # Построение графика с настройками для отступов, шрифтов и легенды
    fig = go.Figure()
    fig.add_trace(go.Scatter(
        x=merged_df['Year'], y=merged_df['Priority'], mode='lines+markers', name='Патентные семейства',
        line=dict(color='rgba(102, 153, 204, 0.8)', width=6, shape='spline'),
        hovertemplate='<b>Год:</b> %{x}<br><b>Количество:</b> %{y}<extra></extra>'
    ))
    fig.add_trace(go.Scatter(
        x=merged_df['Year'], y=merged_df['Publication'], mode='lines+markers', name='Патентные публикации',
        line=dict(color='rgba(153, 204, 102, 0.8)', width=6, shape='spline'),
        hovertemplate='<b>Год:</b> %{x}<br><b>Количество:</b> %{y}<extra></extra>'
    ))
    fig.add_trace(go.Scatter(
        x=merged_df['Year'], y=merged_df['Grant'], mode='lines+markers', name='Патенты',
        line=dict(color='rgba(255, 153, 102, 0.8)', width=6, shape='spline'),
        hovertemplate='<b>Год:</b> %{x}<br><b>Количество:</b> %{y}<extra></extra>'
    ))

    # Настройки макета с увеличенным шрифтом, отступами и улучшенной легендой
    fig.update_layout(
        title='Динамика патентования компании',
        xaxis=dict(
            title='Год',
            tickmode='linear',
            dtick=1,
            range=[2014, 2023],
            title_standoff=25,  # Отступ для заголовка оси X
            tickfont=dict(size=29),  # Увеличенный шрифт для числовых меток оси X (годы)
            automargin=True
        ),
        yaxis=dict(
            title='Количество',
            title_standoff=25,  # Отступ для заголовка оси Y
            tickfont=dict(size=29),  # Увеличенный шрифт для числовых меток оси Y (количество)
            automargin=True
        ),
        template='plotly_white',
        font=dict(family='Arial, sans-serif', size=28, color='black'),  # Единый шрифт
        legend=dict(
            title=dict(text='Категории<br>', font=dict(size=28)),  # Дополнительный отступ для заголовка
            font=dict(family='Arial, sans-serif', size=28, color='black'),  # Шрифт для легенды
            orientation='v',
            y=0.6,
            x=1.05,
            xanchor='left',
            itemclick='toggleothers',  # Позволяет убирать и добавлять линии по клику
            itemsizing='constant',
            traceorder='normal',
            itemwidth=50  # Увеличение расстояния между элементами в легенде
        ),
        margin=dict(t=80, b=40, l=80, r=20)  # Дополнительные отступы
    )

    # Сохранение графика в HTML
    fig.write_html(output_file)
    print("Plot saved as 'trends_by_year_custom_font.html'")


def create_top_patent_word_report_with_translation(file_path_1, file_path_2, output_word_file):
    """
    Загружает данные из двух файлов Excel, объединяет, переводит Title на русский для топ-5 записей,
    сортирует по 'Patent Strength' и сохраняет их в новый Word файл в нужной последовательности.
    """
    
    def clean_text(value):
        """Очищает текст, заменяя начальные символы новой строки (\n) на точку с запятой (;)."""
        if isinstance(value, str):
            return value.replace('\n', ';\n').strip()
        return value

    # Инициализация переводчика
    translator = Translator()

    # Загрузка данных из первого файла
    df1 = pd.read_excel(file_path_1)
    df2 = pd.read_excel(file_path_2)

    # Переименовать нужные столбцы в первом файле
    columns_mapping = {'Название на английском языке': 'Title', 'Сила патентного семейства': 'Patent strength'}
    df1_selected = df1[list(columns_mapping.keys())].rename(columns=columns_mapping)

    # Выбираем нужные столбцы из второго файла
    additional_columns = {'Current assignees': 'Current Assignees', 'Earliest publication number': 'Earliest Publication Number'}
    df2_selected = df2[list(additional_columns.keys())].rename(columns=additional_columns)

    # Объединяем данные из двух файлов
    combined_df = pd.concat([df2_selected, df1_selected], axis=1)

    # Очистка строк от начальных \n
    for col in combined_df.columns:
        combined_df[col] = combined_df[col].apply(clean_text)

    # Сортировка по "Patent strength" и выбор топ-5 записей
    top_5_df = combined_df.sort_values(by='Patent strength', ascending=False).head(5)

    # Перевод столбца Title на русский для топ-5
    top_5_df['Title (Russian)'] = top_5_df['Title'].apply(lambda x: translator.translate(x, src='en', dest='ru').text)

    # Переименование и установка нужного порядка столбцов
    top_5_df.insert(0, '№', range(1, 6))
    top_5_df = top_5_df.rename(columns={
        'Earliest Publication Number': 'Номер публикации',
        'Current Assignees': 'Название компании',
        'Title': 'Название изобретения на английском',
        'Title (Russian)': 'Название изобретения на русском',
        'Patent strength': 'Сила патента'
    })
    top_5_df = top_5_df[['№', 'Номер публикации', 'Название компании', 'Название изобретения на английском', 'Название изобретения на русском', 'Сила патента']]

    # Создание документа Word
    doc = Document()
    doc.add_heading('Топ-5 патентов по силе патента', level=1)

    # Добавление таблицы
    table = doc.add_table(rows=1, cols=len(top_5_df.columns))
    table.style = 'Table Grid'

    # Заполнение заголовков таблицы
    hdr_cells = table.rows[0].cells
    for i, column_name in enumerate(top_5_df.columns):
        hdr_cells[i].text = column_name

    # Заполнение данных таблицы
    for index, row in top_5_df.iterrows():
        row_cells = table.add_row().cells
        for i, item in enumerate(row):
            row_cells[i].text = str(item)

    # Форматирование шрифта для таблицы
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True  # Заголовок таблицы полужирный
                run.font.size = Pt(12)  # Размер шрифта заголовка

    # Сохранение документа Word
    doc.save(output_word_file)
    print(f"Топ-5 записей по Patent Strength сохранены в Word файл '{output_word_file}'")



def country_table():
    # Путь к PDF файлу
    pdf_path = 'C:/website_orbit/website_rup/Lockheed/bunch/country_codes.pdf'

    # Извлечение текста из PDF
    with open(pdf_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ''
        for page in pdf_reader.pages:
            text += page.extract_text()

    # Регулярное выражение для извлечения названий стран и их кодов
    country_pattern = re.compile(r'([A-Z][A-Z\s\(\)]+?)\s+\.+\s+([A-Z]{2})')
    country_mapping = {match[1]: match[0].title().strip() for match in country_pattern.findall(text)}

    # Преобразование словаря в DataFrame
    country_df = pd.DataFrame(list(country_mapping.items()), columns=['Country Code', 'Country Name'])

    # Инициализация переводчика
    translator = Translator()

    # Перевод названий стран на русский
    country_df['Country Name Russian'] = country_df['Country Name'].apply(lambda x: translator.translate(x, src='en', dest='ru').text)

    # Сохранение в Excel
    output_path = 'C:/website_orbit/website_rup/Lockheed/bunch/cleaned_country_mapping_with_russian.xlsx'
    country_df.to_excel(output_path, index=False)
    print(f"Файл сохранен по пути: {output_path}")

def country_rank(patent_data_path, country_mapping_path, output_path):

    # Шаг 1: Загрузка данных
    df_patents = pd.read_excel(patent_data_path)
    df_country_mapping = pd.read_excel(country_mapping_path)

    # Создание словаря для преобразования кодов стран в полные названия на русском
    country_dict = dict(zip(df_country_mapping['Country Code'], df_country_mapping['True Russian Country Name']))

    # Шаг 2: Извлечение и обработка столбца "Страны (юрисдикции) патентования"
    countries = df_patents['Страны (юрисдикции) патентования'].dropna().str.split(', ')
    countries = countries.explode()
    filtered_countries = countries[~countries.isin(['EP', 'WO'])]

    # Преобразуем коды стран в полные названия с помощью словаря
    filtered_countries = filtered_countries.map(country_dict).dropna()

    # Шаг 3: Подсчет количества упоминаний каждой страны и получение топ-20
    country_counts = filtered_countries.value_counts().head(20).reset_index()
    country_counts.columns = ['Страна (юрисдикция)', 'Число патентных семейств']

    # Шаг 4: Визуализация с единым цветом для каждого столбца
    fig = px.bar(
        country_counts,
        x='Страна (юрисдикция)',
        y='Число патентных семейств',
        title='Рейтинг стран (юрисдикций)',
    )
    fig.update_traces(
        marker_color='royalblue',
        text=country_counts['Число патентных семейств'],
        textposition='outside'
    )
    fig.update_layout(
        xaxis_title='Страна (юрисдикция)',
        yaxis_title='Число патентных семейств',
        font=dict(family='Arial, sans-serif', size=26, color='black'),
        xaxis=dict(
            title_standoff=25,  # Отступ для заголовка оси X
            tickfont=dict(size=29),  # Шрифт для меток на оси X
            automargin=True
        ),
        yaxis=dict(
            title_standoff=25,  # Отступ для заголовка оси Y
            tickfont=dict(size=29),  # Шрифт для меток на оси Y
            automargin=True
        ),
        legend=dict(
            title='Категории',
            font=dict(family='Arial, sans-serif', size=26, color='black'),
            orientation='v',
            y=0.6,
            x=1.05,
            xanchor='left'
        ),
        margin=dict(t=80, b=40, l=80, r=40)
    )


    fig.write_html(output_path)
    print(f"График сохранен в {output_path}")

def inventors_rank(patent_data_path, output_path):


    # Загрузка данных
    df_patents = pd.read_excel(patent_data_path)

    # Столбец с авторами
    inventor_column = 'Inventors'
    df_patents[inventor_column] = df_patents[inventor_column].dropna().str.split('\n')
    df_exploded = df_patents.explode(inventor_column)

    # Подсчет числа патентных семейств для каждого автора
    df_exploded['Число патентных семейств'] = 1
    inventor_counts = df_exploded.groupby(inventor_column).count()['Число патентных семейств'].reset_index()
    top_10_inventors = inventor_counts.nlargest(10, 'Число патентных семейств')

    # Визуализация с единым цветом для каждого столбца
    fig = px.bar(
        top_10_inventors,
        x='Inventors',
        y='Число патентных семейств',
        title='Рейтинг авторов изобретений',
        labels={'Inventors': 'Автор изобретения', 'Число патентных семейств': 'Число патентных семейств'},
    )
    fig.update_traces(
        marker_color='darkorange',
        text=top_10_inventors['Число патентных семейств'],
        textposition='outside'
    )
    fig.update_layout(
        xaxis_title='Автор изобретения',
        yaxis_title='Число патентных семейств',
        font=dict(family='Arial, sans-serif', size=26, color='black'),
        xaxis=dict(
            title_standoff=25,  # Отступ для заголовка оси X
            tickfont=dict(size=29),  # Шрифт для меток на оси X
            automargin=True
        ),
        yaxis=dict(
            title_standoff=25,  # Отступ для заголовка оси Y
            tickfont=dict(size=29),  # Шрифт для меток на оси Y
            automargin=True
        ),
        legend=dict(
            title='Категории',
            font=dict(family='Arial, sans-serif', size=26, color='black'),
            orientation='v',
            y=0.6,
            x=1.05,
            xanchor='left'
        ),
        margin=dict(t=80, b=40, l=80, r=40)
    )

    fig.write_html(output_path)
    print(f"График сохранен в {output_path}")


def create_legal_status_pie_charts(file_path, output_dir):
    """
    Загружает данные, очищает их, создает и сохраняет два круговых графика для столбцов 
    "Legal state (Alive, Dead)" и "Legal status (Pending, Granted, Revoked, Expired, Lapsed)".
    """

    # Загрузка данных из Excel
    df = pd.read_excel(file_path)

    # Внутренняя функция для обработки статусов и подсчета значений
    def process_legal_status(df, column_name, values):
        # Извлекаем только значения статусов (например, Alive, Dead, Lapsed и т.д.)
        status_counts = df[column_name].dropna().apply(lambda x: re.findall(r'\b(?:' + '|'.join(values) + r')\b', x))
        status_counts = status_counts.explode().value_counts().reset_index()
        status_counts.columns = ['Status', 'Count']
        return status_counts

    # Внутренняя функция для создания кругового графика с настройкой цветов и отступом для легенды
    def plot_pie_chart(data, title, output_path, status_labels, colors):
        data['Status'] = data['Status'].replace(status_labels)
        fig = px.pie(data, names='Status', values='Count', title=title)

        fig.update_traces(
            textinfo='percent+label', 
            textposition='outside',
            marker=dict(colors=colors)
        )
        
        # Настройки для размещения легенды точно в правом нижнем углу
        fig.update_layout(
            font=dict(family='Arial, sans-serif', size=26, color='black'),
            legend=dict(
                title_text='Категории',
                orientation="v",  # Вертикальное расположение легенды
                x=100000,  # Крайний правый угол
                y=11110,  # Крайний нижний угол
                xanchor="right",
                yanchor="bottom",
                font=dict(size=20),
                traceorder="normal"
            )
        )
        
        fig.write_html(output_path)
        print(f"Круговой график '{title}' сохранен как '{output_path}'")

    # Обработка и создание графика для Legal state (Alive, Dead)
    legal_state_values = ['ALIVE', 'DEAD']
    legal_state_data = process_legal_status(df, 'Legal state (Alive, Dead)', legal_state_values)
    legal_state_labels = {'ALIVE': 'Действующие патентные семейства', 'DEAD': 'Недействующие патентные семейства'}
    legal_state_colors = ['red', '#4682B4']  # Цвета: ALIVE - светло-оранжевый, DEAD - синий
    plot_pie_chart(
        legal_state_data, 
        'Правовое состояние с учётом делопроизводства', 
        f"{output_dir}/legal_state_distribution.html",
        legal_state_labels,
        legal_state_colors
    )

    # Обработка и создание графика для Legal status (Pending, Granted, Revoked, Expired, Lapsed)
    legal_status_values = ['PENDING', 'GRANTED', 'REVOKED', 'EXPIRED', 'LAPSED']
    legal_status_data = process_legal_status(df, 'Legal status (Pending, Granted, Revoked, Expired, Lapsed)', legal_status_values)
    legal_status_labels = {
        'PENDING': 'Заявки на рассмотрении',
        'GRANTED': 'Действующие патенты',
        'REVOKED': 'Отозванные заявки',
        'EXPIRED': 'Патенты с истекшим сроком действия',
        'LAPSED': 'Патенты, прекратившие действие по иным причинам'
    }
    legal_status_colors = ['#B3D7D4', '#4682B4', '#8A2BE2', '#32CD32', 'red']  # Цвета для каждой категории
    plot_pie_chart(
        legal_status_data, 
        'Правовой статус с учётом делопроизводства', 
        f"{output_dir}/legal_status_distribution.html",
        legal_status_labels,
        legal_status_colors
    )

company = 'Leonardo'
yaml_file = '/Users/igorkomissarov/ProjectOffice_FIPS Dropbox/Игорь Комиссаров/WorkPlace/bunch/column_names.yaml'
bunch_file = '/Users/igorkomissarov/ProjectOffice_FIPS Dropbox/Игорь Комиссаров/WorkPlace/bunch/' + company + '/Diversity bunch ' + company + '.xlsx'
output_dir = '/Users/igorkomissarov/ProjectOffice_FIPS Dropbox/Игорь Комиссаров/WorkPlace/bunch/' + company
country_mapping_path = '/Users/igorkomissarov/ProjectOffice_FIPS Dropbox/Игорь Комиссаров/WorkPlace/bunch/Разное/Расшифровка двухбуквенных кодов юрисдикций.xlsx'
file_path_1 = '/Users/igorkomissarov/ProjectOffice_FIPS Dropbox/Игорь Комиссаров/WorkPlace/bunch/' + company + '/Restore_' + company + '.xlsx'

grapth(yaml_file, bunch_file, output_dir + '/graph_' + company + '.html')

create_top_patent_word_report_with_translation(file_path_1, bunch_file, output_dir + '/top_5_combined_patents_by_strength_with_translation.docx')

country_rank(file_path_1, country_mapping_path, output_dir + '/country_rank_' + company + '.html')

inventors_rank(bunch_file, output_dir + '/top_inventor_' + company + '.html')

create_legal_status_pie_charts(bunch_file, output_dir)
